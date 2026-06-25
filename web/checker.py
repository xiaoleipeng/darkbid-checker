"""
Word暗标编制要求检查工具 - Pyodide版核心逻辑
在浏览器中通过Pyodide运行，提供 check_fix_annotate() 函数。
"""
import io
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# 常量
FONT_NAME = '宋体'
FONT_SIZE = Pt(14)
PAGE_WIDTH = Cm(21.0)
PAGE_HEIGHT = Cm(29.7)
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.0)
MARGIN_LEFT = Cm(2.0)
MARGIN_RIGHT = Cm(2.0)
LINE_SPACING_VAL = Pt(30)
BLACK_COLOR = RGBColor(0, 0, 0)
AUTHOR = "暗标检查工具"
INITIALS = "AH"

PUNCTUATION_MAP = {
    ',': '，', '.': '。', ':': '：', ';': '；',
    '!': '！', '?': '？', '(': '（', ')': '）',
    '[': '【', ']': '】',
}
PUNCT_SET = set(PUNCTUATION_MAP.keys())
DIGIT_PUNCT_PATTERN = re.compile(r'[a-zA-Z0-9]+([.,][a-zA-Z0-9]+)+')


def _clean_alnum_punct(text):
    """删除字母数字间用.,连接的整段序列（如1.3.5.8、report.docx、www.example.com）"""
    return DIGIT_PUNCT_PATTERN.sub('', text)

# 预计算 qn 值避免重复调用
QN_WPR = qn('w:rPr')
QN_WFONTS = qn('w:rFonts')
QN_EASTASIA = qn('w:eastAsia')
QN_INSTR = qn('w:instrText')
QN_TC = qn('w:tc')
QN_TBL = qn('w:tbl')

# 批注合并阈值（不再限制）
# details 最大条目数（不再限制）


def _set_font_name(run):
    run.font.name = FONT_NAME
    rpr = run._element.find(QN_WPR)
    if rpr is None:
        rpr = run._element.makeelement(QN_WPR, {})
        run._element.insert(0, rpr)
    rFonts = rpr.find(QN_WFONTS)
    if rFonts is None:
        rFonts = rpr.makeelement(QN_WFONTS, {})
        rpr.insert(0, rFonts)
    rFonts.set(QN_EASTASIA, FONT_NAME)


def _fix_punctuation_in_run(run):
    text = run.text
    if not any(ch in PUNCT_SET for ch in text):
        return
    result = []
    for i, ch in enumerate(text):
        if ch in PUNCTUATION_MAP:
            if ch in '.,':
                # 前后有字母或数字时保留（数字小数、版本号、文件名、网址、带单位数字等）
                before_alnum = i > 0 and (text[i-1].isdigit() or text[i-1].isalpha())
                after_alnum = i < len(text)-1 and (text[i+1].isdigit() or text[i+1].isalpha())
                if before_alnum and after_alnum:
                    result.append(ch)
                    continue
            result.append(PUNCTUATION_MAP[ch])
        else:
            result.append(ch)
    run.text = ''.join(result)


_RULE_NAMES = {
    WD_LINE_SPACING.SINGLE: "单倍行距",
    WD_LINE_SPACING.ONE_POINT_FIVE: "1.5倍行距",
    WD_LINE_SPACING.DOUBLE: "2倍行距",
    WD_LINE_SPACING.MULTIPLE: "多倍行距",
    WD_LINE_SPACING.AT_LEAST: "最小值",
    WD_LINE_SPACING.EXACTLY: "固定值",
}


def _describe_line_spacing(rule, spacing):
    """根据 rule 和 spacing 值生成可读描述"""
    if rule is None and spacing is None:
        return "单倍行距"
    if rule == WD_LINE_SPACING.MULTIPLE and spacing:
        return f"{spacing:.2f}倍行距"
    if rule == WD_LINE_SPACING.AT_LEAST and spacing:
        return f"最小值{spacing.pt:.1f}磅"
    if rule == WD_LINE_SPACING.EXACTLY and spacing:
        return f"固定值{spacing.pt:.1f}磅"
    return _RULE_NAMES.get(rule, "单倍行距")


def _get_actual_line_spacing_desc(para):
    """获取段落实际的行间距描述（含样式继承）"""
    pf = para.paragraph_format
    rule = pf.line_spacing_rule
    spacing = pf.line_spacing
    # 如果段落本身有设置，直接用
    if rule is not None:
        return _describe_line_spacing(rule, spacing)
    # 从样式继承
    style = para.style
    if style and style.paragraph_format:
        sr = style.paragraph_format.line_spacing_rule
        ss = style.paragraph_format.line_spacing
        if sr is not None:
            return _describe_line_spacing(sr, ss)
    return "单倍行距"


def _process_runs_font_color(runs, in_table, categories):
    """批量处理 runs 的字体和颜色，返回 (fix_count, fix_descs)"""
    check_font = 'font' in categories
    check_color = 'color' in categories
    fix_count = 0
    fix_descs = []

    for run in runs:
        if not run.text.strip():
            continue
        font = run.font

        if check_font:
            if not in_table:
                # 字体名
                if font.name and font.name != FONT_NAME:
                    fix_descs.append(f"字体: {font.name} → 宋体")
                    _set_font_name(run)
                    fix_count += 1
                else:
                    rpr = run._element.find(QN_WPR)
                    if rpr is not None:
                        rFonts = rpr.find(QN_WFONTS)
                        if rFonts is not None:
                            ea = rFonts.get(QN_EASTASIA)
                            if ea and ea != FONT_NAME:
                                fix_descs.append(f"中文字体: {ea} → 宋体")
                                _set_font_name(run)
                                fix_count += 1
                # 字号
                if font.size and font.size != FONT_SIZE:
                    fix_descs.append(f"字号: {font.size.pt:.1f}pt → 14pt(四号)")
                    font.size = FONT_SIZE
                    fix_count += 1
            # 加粗/倾斜/下划线（表格内外均检查）
            if font.bold:
                fix_descs.append("去除加粗")
                font.bold = False
                fix_count += 1
            if font.italic:
                fix_descs.append("去除倾斜")
                font.italic = False
                fix_count += 1
            if font.underline:
                fix_descs.append("去除下划线")
                font.underline = False
                fix_count += 1

        if check_color:
            if font.color and font.color.rgb and font.color.rgb != BLACK_COLOR:
                fix_descs.append(f"颜色: #{font.color.rgb} → 黑色")
                font.color.rgb = BLACK_COLOR
                fix_count += 1

    return fix_count, fix_descs


def check_fix_annotate(doc_path, output_path, keyword_map=None, categories=None, add_comments=True, progress_callback=None):
    """核心修复逻辑，返回 (details, fix_count, warn_count)"""
    ALL_CATEGORIES = {'page', 'spacing', 'align', 'font', 'color', 'punct', 'identity'}
    if categories is None:
        categories = ALL_CATEGORIES

    # 支持路径或BytesIO
    if isinstance(doc_path, str):
        with open(doc_path, 'rb') as f:
            doc = Document(io.BytesIO(f.read()))
    else:
        doc = Document(doc_path)

    fix_count = 0
    warn_count = 0
    details = []

    # 判断段落是否在表格内（通过XML父元素，不用id()）
    def _is_in_table(para):
        parent = para._element.getparent()
        while parent is not None:
            if parent.tag.endswith('}tc') or parent.tag.endswith('}tbl'):
                return True
            parent = parent.getparent()
        return False

    first_para = doc.paragraphs[0] if doc.paragraphs else None

    # TOC目录检测 + 页面设置（合并到一次遍历前处理）
    if 'page' in categories:
        # TOC 检测只扫前50段（目录通常在文档开头）
        for para in doc.paragraphs[:50]:
            for instr in para._element.findall('.//' + QN_INSTR):
                if instr.text and 'TOC' in instr.text.upper():
                    details.append(("目录", "⚠️ 存在自动目录，需手动删除"))
                    warn_count += 1
                    break

        # 页面设置
        page_fixes = []
        tolerance = Cm(0.1)
        margin_tolerance = Cm(0.05)
        for section in doc.sections:
            if section.page_width and abs(section.page_width - PAGE_WIDTH) > tolerance:
                page_fixes.append(f"纸张宽度 → 21.0cm")
                section.page_width = PAGE_WIDTH
            if section.page_height and abs(section.page_height - PAGE_HEIGHT) > tolerance:
                page_fixes.append(f"纸张高度 → 29.7cm")
                section.page_height = PAGE_HEIGHT
            if section.top_margin is not None and abs(section.top_margin - MARGIN_TOP) > margin_tolerance:
                page_fixes.append(f"上边距 → 2.5cm")
                section.top_margin = MARGIN_TOP
            if section.bottom_margin is not None and abs(section.bottom_margin - MARGIN_BOTTOM) > margin_tolerance:
                page_fixes.append(f"下边距 → 2.0cm")
                section.bottom_margin = MARGIN_BOTTOM
            if section.left_margin is not None and abs(section.left_margin - MARGIN_LEFT) > margin_tolerance:
                page_fixes.append(f"左边距 → 2.0cm")
                section.left_margin = MARGIN_LEFT
            if section.right_margin is not None and abs(section.right_margin - MARGIN_RIGHT) > margin_tolerance:
                page_fixes.append(f"右边距 → 2.0cm")
                section.right_margin = MARGIN_RIGHT
            if section.header and section.header.is_linked_to_previous is False:
                if ''.join(p.text for p in section.header.paragraphs).strip():
                    page_fixes.append("已清除页眉")
                    for p in section.header.paragraphs:
                        p.clear()
            if section.footer and section.footer.is_linked_to_previous is False:
                if ''.join(p.text for p in section.footer.paragraphs).strip():
                    page_fixes.append("已清除页脚/页码")
                    for p in section.footer.paragraphs:
                        p.clear()
        if page_fixes:
            fix_count += len(page_fixes)
            for f in page_fixes:
                details.append(("页面设置", f))

    # 页数估算
    total_paras = len(doc.paragraphs)
    estimated_pages = total_paras // 25
    if estimated_pages > 300:
        details.append(("全文", f"⚠️ 估算页数约{estimated_pages}页，可能超过300页限制"))
        warn_count += 1

    # 预编译关键词检查
    kw_items = list(keyword_map.items()) if keyword_map and 'identity' in categories else []

    check_spacing = 'spacing' in categories
    check_align = 'align' in categories
    check_punct = 'punct' in categories
    check_font_or_color = 'font' in categories or 'color' in categories

    # 主循环：逐段处理（跳过表格内段落，后面单独处理）
    comment_count = 0
    total_paras_list = doc.paragraphs
    total_para_count = len(total_paras_list)
    for para_idx, para in enumerate(total_paras_list):
        if progress_callback and para_idx % 10 == 0:
            pct = 10 + 70 * para_idx / max(total_para_count, 1)
            progress_callback(pct, f"正在处理第 {para_idx}/{total_para_count} 段...")
        text = para.text
        if not text.strip() or not para.runs:
            continue

        in_table = _is_in_table(para)
        if in_table:
            continue  # 表格内段落在下面统一处理

        para_descs = []

        if check_spacing:
            pf = para.paragraph_format
            # 获取实际生效的行间距（含样式继承）
            eff_rule = pf.line_spacing_rule
            eff_spacing = pf.line_spacing
            if eff_rule is None:
                style = para.style
                if style and style.paragraph_format:
                    eff_rule = style.paragraph_format.line_spacing_rule
                    eff_spacing = style.paragraph_format.line_spacing
            need_fix = False
            if eff_rule != WD_LINE_SPACING.EXACTLY:
                need_fix = True
            elif eff_spacing is None or abs(eff_spacing - LINE_SPACING_VAL) > Pt(0.5):
                need_fix = True
            if need_fix:
                old_desc = _get_actual_line_spacing_desc(para)
                para_descs.append(f"行间距: {old_desc} → 固定值30磅")
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = LINE_SPACING_VAL
                fix_count += 1
            if pf.space_before and pf.space_before > Pt(0):
                para_descs.append(f"段前: {pf.space_before.pt:.1f}磅 → 0")
                pf.space_before = Pt(0)
                fix_count += 1
            if pf.space_after and pf.space_after > Pt(0):
                para_descs.append(f"段后: {pf.space_after.pt:.1f}磅 → 0")
                pf.space_after = Pt(0)
                fix_count += 1

        if check_align:
            pf = para.paragraph_format
            if pf.alignment is not None and pf.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                align_names = {WD_ALIGN_PARAGRAPH.CENTER: "居中", WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
                               WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐", WD_ALIGN_PARAGRAPH.DISTRIBUTE: "分散对齐"}
                para_descs.append(f"对齐: {align_names.get(pf.alignment, '?')} → 左对齐")
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                fix_count += 1
            if pf.first_line_indent is None or abs(pf.first_line_indent - Pt(28)) > Pt(1):
                old_indent = f"{pf.first_line_indent.pt:.1f}pt" if pf.first_line_indent else "无"
                para_descs.append(f"首行缩进: {old_indent} → 2字符")
                pf.first_line_indent = Pt(28)
                fix_count += 1
            if text[0] in (' ', '\u3000'):
                para_descs.append("删除开头空格")
                para.runs[0].text = para.runs[0].text.lstrip(' \u3000')
                fix_count += 1

        if check_punct:
            if PUNCT_SET & set(text):
                cleaned = _clean_alnum_punct(text)
                en_puncts = [ch for ch in cleaned if ch in PUNCT_SET]
                if en_puncts:
                    samples = list(set(en_puncts))[:5]
                    mapping = ' '.join(f"'{c}'→'{PUNCTUATION_MAP[c]}'" for c in samples)
                    para_descs.append(f"标点: {mapping}")
                    for run in para.runs:
                        _fix_punctuation_in_run(run)
                    fix_count += 1

        if kw_items:
            for kw, repl in kw_items:
                if kw in text:
                    para_descs.append(f"'{kw}' → '{repl}'")
                    for run in para.runs:
                        if kw in run.text:
                            run.text = run.text.replace(kw, repl)
                    fix_count += 1
                    warn_count += 1

        if check_font_or_color:
            fc, font_descs = _process_runs_font_color(para.runs, False, categories)
            fix_count += fc
            para_descs.extend(font_descs)

        # 加批注（带修改明细）
        if para_descs and add_comments:
            doc.add_comment(
                runs=para.runs,
                text="【已修复】\n" + "\n".join(f"• {d}" for d in para_descs),
                author=AUTHOR, initials=INITIALS
            )
            comment_count += 1

        if para_descs:
            para_loc = text[:25] + "..." if len(text) > 25 else text
            for d in para_descs:
                details.append((para_loc, d))

    # 表格内段落处理
    if check_font_or_color or check_align:
        if progress_callback:
            progress_callback(82, "正在处理表格内容...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text.strip() or not para.runs:
                            continue
                        tbl_descs = []

                        if check_align:
                            pf = para.paragraph_format
                            if pf.alignment is not None and pf.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                                align_names = {WD_ALIGN_PARAGRAPH.CENTER: "居中", WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
                                               WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐"}
                                tbl_descs.append(f"对齐: {align_names.get(pf.alignment, '?')} → 左对齐")
                                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                fix_count += 1

                        if check_font_or_color:
                            fc, font_descs = _process_runs_font_color(para.runs, True, categories)
                            fix_count += fc
                            tbl_descs.extend(font_descs)

                        if tbl_descs and add_comments:
                            doc.add_comment(
                                runs=para.runs,
                                text="【已修复-表格内】\n" + "\n".join(f"• {d}" for d in tbl_descs),
                                author=AUTHOR, initials=INITIALS
                            )
                            comment_count += 1

                        if tbl_descs:
                            para_loc = "[表格内] " + (para.text[:15] + "..." if len(para.text) > 15 else para.text)
                            for d in tbl_descs:
                                details.append((para_loc, d))

    if fix_count == 0 and warn_count == 0:
        return details, 0, 0

    if progress_callback:
        progress_callback(92, "正在保存文件...")

    if isinstance(output_path, str):
        buf = io.BytesIO()
        doc.save(buf)
        with open(output_path, 'wb') as f:
            f.write(buf.getvalue())
    else:
        doc.save(output_path)
    return details, fix_count, warn_count
