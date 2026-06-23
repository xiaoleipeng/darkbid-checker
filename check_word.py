#!/usr/bin/env python3
"""
Word文档暗标编制要求检查工具
检查技术标文档是否符合暗标编制规范，标注不符合项，提供交互式确认修改功能。
"""

import sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


# 常量定义
FONT_NAME = '宋体'
FONT_SIZE = Pt(14)  # 四号字 = 14pt
PAGE_WIDTH = Cm(21.0)  # A4宽度
PAGE_HEIGHT = Cm(29.7)  # A4高度
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.0)
MARGIN_LEFT = Cm(2.0)
MARGIN_RIGHT = Cm(2.0)
LINE_SPACING = Pt(30)  # 固定值30磅
MAX_PAGES = 300
BLACK_COLOR = RGBColor(0, 0, 0)


class Issue:
    """表示一个不符合项"""
    def __init__(self, description, location, fix_func=None, category=""):
        self.description = description
        self.location = location
        self.fix_func = fix_func  # 修复函数
        self.category = category
        self.selected = True  # 默认选中修复


def check_page_setup(doc):
    """检查页面设置：A4纸张、页边距"""
    issues = []
    for i, section in enumerate(doc.sections):
        loc = f"第{i+1}节"
        # 检查纸张大小（允许误差）
        tolerance = Cm(0.1)
        if section.page_width and abs(section.page_width - PAGE_WIDTH) > tolerance:
            issues.append(Issue(
                f"纸张宽度不是A4（当前: {section.page_width.cm:.1f}cm，要求: 21.0cm）",
                loc, lambda s=section: setattr(s, 'page_width', PAGE_WIDTH), "版面要求"))
        if section.page_height and abs(section.page_height - PAGE_HEIGHT) > tolerance:
            issues.append(Issue(
                f"纸张高度不是A4（当前: {section.page_height.cm:.1f}cm，要求: 29.7cm）",
                loc, lambda s=section: setattr(s, 'page_height', PAGE_HEIGHT), "版面要求"))
        # 检查页边距
        margin_tolerance = Cm(0.05)
        if section.top_margin is not None and abs(section.top_margin - MARGIN_TOP) > margin_tolerance:
            issues.append(Issue(
                f"上边距不符（当前: {section.top_margin.cm:.2f}cm，要求: 2.5cm）",
                loc, lambda s=section: setattr(s, 'top_margin', MARGIN_TOP), "排版要求"))
        if section.bottom_margin is not None and abs(section.bottom_margin - MARGIN_BOTTOM) > margin_tolerance:
            issues.append(Issue(
                f"下边距不符（当前: {section.bottom_margin.cm:.2f}cm，要求: 2.0cm）",
                loc, lambda s=section: setattr(s, 'bottom_margin', MARGIN_BOTTOM), "排版要求"))
        if section.left_margin is not None and abs(section.left_margin - MARGIN_LEFT) > margin_tolerance:
            issues.append(Issue(
                f"左边距不符（当前: {section.left_margin.cm:.2f}cm，要求: 2.0cm）",
                loc, lambda s=section: setattr(s, 'left_margin', MARGIN_LEFT), "排版要求"))
        if section.right_margin is not None and abs(section.right_margin - MARGIN_RIGHT) > margin_tolerance:
            issues.append(Issue(
                f"右边距不符（当前: {section.right_margin.cm:.2f}cm，要求: 2.0cm）",
                loc, lambda s=section: setattr(s, 'right_margin', MARGIN_RIGHT), "排版要求"))
        # 检查页眉页脚
        if section.header and section.header.is_linked_to_previous is False:
            header_text = ''.join(p.text for p in section.header.paragraphs)
            if header_text.strip():
                issues.append(Issue(
                    f"存在页眉内容: '{header_text.strip()[:20]}...'",
                    loc, lambda s=section: _clear_header(s), "排版要求"))
        if section.footer and section.footer.is_linked_to_previous is False:
            footer_text = ''.join(p.text for p in section.footer.paragraphs)
            if footer_text.strip():
                issues.append(Issue(
                    f"存在页脚内容: '{footer_text.strip()[:20]}...'",
                    loc, lambda s=section: _clear_footer(s), "排版要求"))
        # 检查是否有页码（different_first_page_header_footer 或 页脚中的PAGE字段）
        if section.footer:
            for para in section.footer.paragraphs:
                for run in para.runs:
                    if 'PAGE' in run.text or any(
                        fld.get(qn('w:instr'), '').strip().upper().startswith('PAGE')
                        for fld in run._element.findall(qn('w:fldChar'))
                    ):
                        issues.append(Issue(
                            "存在页码",
                            loc, lambda s=section: _clear_footer(s), "排版要求"))
                        break
    return issues


def _clear_header(section):
    """清除页眉"""
    for para in section.header.paragraphs:
        para.clear()


def _clear_footer(section):
    """清除页脚"""
    for para in section.footer.paragraphs:
        para.clear()


def _is_in_table(paragraph):
    """判断段落是否在表格内"""
    parent = paragraph._element.getparent()
    while parent is not None:
        if parent.tag.endswith('}tc') or parent.tag.endswith('}tbl'):
            return True
        parent = parent.getparent()
    return False


def check_paragraph_format(doc):
    """检查段落格式：行间距、对齐方式、首行缩进"""
    issues = []
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        if _is_in_table(para):
            continue
        loc = f"第{i+1}段: '{para.text[:30]}...'" if len(para.text) > 30 else f"第{i+1}段: '{para.text}'"
        pf = para.paragraph_format

        # 检查行间距 - 固定值30磅
        if pf.line_spacing_rule != WD_LINE_SPACING.EXACTLY or pf.line_spacing != LINE_SPACING:
            current = "未设置"
            if pf.line_spacing is not None:
                if pf.line_spacing_rule == WD_LINE_SPACING.MULTIPLE:
                    current = f"{pf.line_spacing:.1f}倍"
                elif isinstance(pf.line_spacing, float):
                    current = f"{pf.line_spacing:.1f}倍"
                else:
                    current = f"固定值{pf.line_spacing.pt:.0f}磅" if pf.line_spacing_rule == WD_LINE_SPACING.EXACTLY else f"{pf.line_spacing.pt:.0f}磅"
            issues.append(Issue(
                f"行间距不符（当前: {current}，要求: 固定值30磅）",
                loc, lambda p=para: _set_line_spacing(p), "排版要求"))

        # 检查对齐方式 - 左对齐
        if pf.alignment is not None and pf.alignment != WD_ALIGN_PARAGRAPH.LEFT:
            align_map = {
                WD_ALIGN_PARAGRAPH.CENTER: "居中",
                WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐",
            }
            current = align_map.get(pf.alignment, str(pf.alignment))
            issues.append(Issue(
                f"对齐方式不符（当前: {current}，要求: 左对齐）",
                loc, lambda p=para: setattr(p.paragraph_format, 'alignment', WD_ALIGN_PARAGRAPH.LEFT), "排版要求"))

        # 检查首行缩进2字符
        if pf.first_line_indent is None or abs(pf.first_line_indent - Pt(28)) > Pt(1):
            # 2字符 ≈ 2 * 14pt = 28pt (四号字)
            current = f"{pf.first_line_indent.pt:.0f}pt" if pf.first_line_indent else "无"
            issues.append(Issue(
                f"首行缩进不符（当前: {current}，要求: 2字符≈28pt）",
                loc, lambda p=para: setattr(p.paragraph_format, 'first_line_indent', Pt(28)), "排版要求"))

        # 检查段前段后间距
        if pf.space_before and pf.space_before > Pt(0):
            issues.append(Issue(
                f"段前有间距（当前: {pf.space_before.pt:.0f}pt，要求: 0）",
                loc, lambda p=para: setattr(p.paragraph_format, 'space_before', Pt(0)), "排版要求"))
        if pf.space_after and pf.space_after > Pt(0):
            issues.append(Issue(
                f"段后有间距（当前: {pf.space_after.pt:.0f}pt，要求: 0）",
                loc, lambda p=para: setattr(p.paragraph_format, 'space_after', Pt(0)), "排版要求"))

    return issues


def _set_line_spacing(para):
    """设置行间距为固定值30磅"""
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = LINE_SPACING


def check_font(doc):
    """检查字体：宋体、四号、常规，无加粗/倾斜/下划线/颜色"""
    issues = []
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        if _is_in_table(para):
            continue
        for j, run in enumerate(para.runs):
            if not run.text.strip():
                continue
            loc = f"第{i+1}段-第{j+1}个文本块: '{run.text[:20]}...'" if len(run.text) > 20 else f"第{i+1}段-第{j+1}个文本块: '{run.text}'"
            font = run.font

            # 检查字体名称
            if font.name and font.name != FONT_NAME:
                issues.append(Issue(
                    f"字体不符（当前: {font.name}，要求: 宋体）",
                    loc, lambda r=run: _set_font_name(r), "字体要求"))

            # 检查中文字体
            rpr = run._element.find(qn('w:rPr'))
            if rpr is not None:
                rFonts = rpr.find(qn('w:rFonts'))
                if rFonts is not None:
                    east_asia = rFonts.get(qn('w:eastAsia'))
                    if east_asia and east_asia != FONT_NAME:
                        issues.append(Issue(
                            f"中文字体不符（当前: {east_asia}，要求: 宋体）",
                            loc, lambda r=run: _set_font_name(r), "字体要求"))

            # 检查字号 - 四号 = 14pt
            if font.size and font.size != FONT_SIZE:
                issues.append(Issue(
                    f"字号不符（当前: {font.size.pt:.1f}pt，要求: 14pt/四号）",
                    loc, lambda r=run: setattr(r.font, 'size', FONT_SIZE), "字体要求"))

            # 检查加粗
            if font.bold:
                issues.append(Issue(
                    "存在加粗",
                    loc, lambda r=run: setattr(r.font, 'bold', False), "字体要求"))

            # 检查倾斜
            if font.italic:
                issues.append(Issue(
                    "存在倾斜",
                    loc, lambda r=run: setattr(r.font, 'italic', False), "字体要求"))

            # 检查下划线
            if font.underline:
                issues.append(Issue(
                    "存在下划线",
                    loc, lambda r=run: setattr(r.font, 'underline', False), "字体要求"))

            # 检查颜色 - 必须为黑色
            if font.color and font.color.rgb and font.color.rgb != BLACK_COLOR:
                issues.append(Issue(
                    f"字体颜色不是黑色（当前: #{font.color.rgb}）",
                    loc, lambda r=run: setattr(r.font.color, 'rgb', BLACK_COLOR), "颜色要求"))

    return issues


def _set_font_name(run):
    """设置字体为宋体"""
    run.font.name = FONT_NAME
    # 设置中文字体
    rpr = run._element.find(qn('w:rPr'))
    if rpr is None:
        rpr = run._element.makeelement(qn('w:rPr'), {})
        run._element.insert(0, rpr)
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)


def check_spaces_in_indent(doc):
    """检查是否用空格代替首行缩进"""
    issues = []
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        if _is_in_table(para):
            continue
        if para.text.startswith(' ') or para.text.startswith('\u3000'):
            loc = f"第{i+1}段: '{para.text[:30]}...'" if len(para.text) > 30 else f"第{i+1}段: '{para.text}'"
            issues.append(Issue(
                "段落开头存在空格（应使用首行缩进，不得有空格）",
                loc, lambda p=para: _remove_leading_spaces(p), "排版要求"))
    return issues


def _remove_leading_spaces(para):
    """去除段落开头的空格"""
    if para.runs:
        para.runs[0].text = para.runs[0].text.lstrip(' \u3000')


def check_toc(doc):
    """检查是否存在目录（TOC字段）"""
    issues = []
    for i, para in enumerate(doc.paragraphs):
        # 检查段落中是否有TOC域代码
        for fld in para._element.findall('.//' + qn('w:fldChar')):
            fld_type = fld.get(qn('w:fldCharType'))
            if fld_type == 'begin':
                # 找到域开始，检查后续instrText
                next_elem = fld.getnext()
                while next_elem is not None:
                    if next_elem.tag == qn('w:instrText'):
                        if 'TOC' in next_elem.text.upper():
                            loc = f"第{i+1}段"
                            issues.append(Issue(
                                "存在自动目录（技术标暗标部分不得设置目录）",
                                loc, None, "封面目录"))
                            return issues
                    if next_elem.tag == qn('w:fldChar'):
                        break
                    next_elem = next_elem.getnext()
        # 也检查复杂域的情况
        for instr in para._element.findall('.//' + qn('w:instrText')):
            if instr.text and 'TOC' in instr.text.upper():
                loc = f"第{i+1}段"
                issues.append(Issue(
                    "存在自动目录（技术标暗标部分不得设置目录）",
                    loc, None, "封面目录"))
                return issues
    return issues


# 英文标点 -> 中文标点映射
PUNCTUATION_MAP = {
    ',': '，', '.': '。', ':': '：', ';': '；',
    '!': '！', '?': '？', '(': '（', ')': '）',
    '[': '【', ']': '】',
}
# 需要排除的场景：数字间的小数点/逗号（如3.14, 1,000）
import re
DIGIT_PUNCT_PATTERN = re.compile(r'\d[.,]\d')


def check_punctuation(doc):
    """检查是否全部使用中文标点"""
    issues = []
    en_puncts = set(PUNCTUATION_MAP.keys())
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        if _is_in_table(para):
            continue
        text = para.text
        # 排除数字中的英文标点（如小数点、千分位逗号）
        cleaned = DIGIT_PUNCT_PATTERN.sub('', text)
        found = set()
        for ch in cleaned:
            if ch in en_puncts:
                found.add(ch)
        if found:
            loc = f"第{i+1}段: '{text[:30]}...'" if len(text) > 30 else f"第{i+1}段: '{text}'"
            found_str = ' '.join(f"'{c}'" for c in sorted(found))
            issues.append(Issue(
                f"存在英文标点: {found_str}（要求全部使用中文标点）",
                loc, lambda p=para, fm=found: _fix_punctuation(p, fm), "字体要求"))
    return issues


def _fix_punctuation(para, found_chars):
    """将段落中的英文标点替换为中文标点"""
    for run in para.runs:
        for en, cn in PUNCTUATION_MAP.items():
            if en in found_chars:
                # 保留数字间的英文标点
                new_text = []
                text = run.text
                i = 0
                while i < len(text):
                    if text[i] == en:
                        # 检查前后是否为数字
                        before_digit = i > 0 and text[i-1].isdigit()
                        after_digit = i < len(text)-1 and text[i+1].isdigit()
                        if before_digit and after_digit and en in '.,':
                            new_text.append(en)
                        else:
                            new_text.append(cn)
                    else:
                        new_text.append(text[i])
                    i += 1
                run.text = ''.join(new_text)


def check_identity_info(doc, keywords=None):
    """检查是否包含投标人身份信息"""
    issues = []
    if not keywords:
        return issues
    full_text = '\n'.join(para.text for para in doc.paragraphs)
    for kw in keywords:
        if kw in full_text:
            # 找到具体位置
            for i, para in enumerate(doc.paragraphs):
                if kw in para.text:
                    loc = f"第{i+1}段: '{para.text[:30]}...'" if len(para.text) > 30 else f"第{i+1}段: '{para.text}'"
                    issues.append(Issue(
                        f"发现可能的身份信息: '{kw}'",
                        loc, lambda p=para, k=kw: _remove_keyword(p, k), "身份信息"))
    return issues


def _remove_keyword(para, keyword):
    """从段落中移除关键词（替换为***）"""
    for run in para.runs:
        if keyword in run.text:
            run.text = run.text.replace(keyword, '***')


def check_page_count(doc):
    """检查页数（近似估算）"""
    issues = []
    total_paras = len(doc.paragraphs)
    estimated_pages = total_paras // 25
    if estimated_pages > MAX_PAGES:
        issues.append(Issue(
            f"估算页数约{estimated_pages}页，可能超过300页限制（请在Word中确认实际页数）",
            "全文", None, "页数要求"))
    return issues


def run_all_checks(doc, identity_keywords=None):
    """执行所有检查"""
    all_issues = []
    all_issues.extend(check_page_setup(doc))
    all_issues.extend(check_toc(doc))
    all_issues.extend(check_paragraph_format(doc))
    all_issues.extend(check_font(doc))
    all_issues.extend(check_punctuation(doc))
    all_issues.extend(check_spaces_in_indent(doc))
    all_issues.extend(check_identity_info(doc, identity_keywords))
    all_issues.extend(check_page_count(doc))
    return all_issues


class CheckerApp:
    """交互式GUI应用"""
    def __init__(self, root):
        self.root = root
        self.root.title("Word暗标编制要求检查工具")
        self.root.geometry("1100x700")
        self.doc = None
        self.doc_path = None
        self.issues = []
        self._build_ui()

    def _build_ui(self):
        # 顶部操作栏
        top_frame = tk.Frame(self.root)
        top_frame.pack(fill=tk.X, padx=10, pady=5)

        tk.Button(top_frame, text="选择Word文件", command=self._open_file).pack(side=tk.LEFT)
        self.file_label = tk.Label(top_frame, text="未选择文件", fg="gray")
        self.file_label.pack(side=tk.LEFT, padx=10)

        tk.Button(top_frame, text="全选", command=self._select_all).pack(side=tk.RIGHT, padx=2)
        tk.Button(top_frame, text="全不选", command=self._deselect_all).pack(side=tk.RIGHT, padx=2)
        tk.Button(top_frame, text="应用选中修改", command=self._apply_fixes, bg="#4CAF50", fg="white").pack(side=tk.RIGHT, padx=5)

        # 关键词输入栏（投标人身份信息检测）
        kw_frame = tk.Frame(self.root)
        kw_frame.pack(fill=tk.X, padx=10, pady=2)
        tk.Label(kw_frame, text="身份关键词（逗号分隔，如公司名/人名）:").pack(side=tk.LEFT)
        self.kw_entry = tk.Entry(kw_frame, width=60)
        self.kw_entry.pack(side=tk.LEFT, padx=5)
        tk.Button(kw_frame, text="重新检查", command=self._run_check).pack(side=tk.LEFT)

        # 统计信息
        self.stats_label = tk.Label(self.root, text="", anchor=tk.W)
        self.stats_label.pack(fill=tk.X, padx=10)

        # 问题列表（带滚动条）
        list_frame = tk.Frame(self.root)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        columns = ("select", "category", "description", "location")
        self.tree = ttk.Treeview(list_frame, columns=columns, show="headings", selectmode="extended")
        self.tree.heading("select", text="✓")
        self.tree.heading("category", text="类别")
        self.tree.heading("description", text="问题描述")
        self.tree.heading("location", text="位置")
        self.tree.column("select", width=30, anchor=tk.CENTER)
        self.tree.column("category", width=80)
        self.tree.column("description", width=500)
        self.tree.column("location", width=400)

        scrollbar_y = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.tree.yview)
        scrollbar_x = ttk.Scrollbar(list_frame, orient=tk.HORIZONTAL, command=self.tree.xview)
        self.tree.configure(yscrollcommand=scrollbar_y.set, xscrollcommand=scrollbar_x.set)

        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
        scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)

        self.tree.bind("<Double-1>", self._toggle_select)

    def _open_file(self):
        path = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")])
        if not path:
            return
        self.doc_path = path
        self.file_label.config(text=Path(path).name, fg="black")
        try:
            self.doc = Document(path)
        except Exception as e:
            messagebox.showerror("错误", f"无法打开文件: {e}")
            return
        self._run_check()

    def _run_check(self):
        if not self.doc:
            return
        # 获取身份关键词
        kw_text = self.kw_entry.get().strip()
        keywords = [k.strip() for k in kw_text.split(',') if k.strip()] if kw_text else None
        self.issues = run_all_checks(self.doc, keywords)
        self._refresh_list()
        fixable = sum(1 for iss in self.issues if iss.fix_func)
        self.stats_label.config(
            text=f"共发现 {len(self.issues)} 个问题，其中 {fixable} 个可自动修复")

    def _refresh_list(self):
        self.tree.delete(*self.tree.get_children())
        for i, issue in enumerate(self.issues):
            check = "✓" if issue.selected and issue.fix_func else ("—" if not issue.fix_func else "")
            self.tree.insert("", tk.END, iid=str(i), values=(check, issue.category, issue.description, issue.location))

    def _toggle_select(self, event):
        item = self.tree.identify_row(event.y)
        if item:
            idx = int(item)
            issue = self.issues[idx]
            if issue.fix_func:
                issue.selected = not issue.selected
                check = "✓" if issue.selected else ""
                self.tree.set(item, "select", check)

    def _select_all(self):
        for issue in self.issues:
            if issue.fix_func:
                issue.selected = True
        self._refresh_list()

    def _deselect_all(self):
        for issue in self.issues:
            issue.selected = False
        self._refresh_list()

    def _apply_fixes(self):
        selected = [iss for iss in self.issues if iss.selected and iss.fix_func]
        if not selected:
            messagebox.showinfo("提示", "未选择任何可修复的问题")
            return
        if not messagebox.askyesno("确认", f"确定要应用 {len(selected)} 项修改？"):
            return
        errors = []
        for issue in selected:
            try:
                issue.fix_func()
            except Exception as e:
                errors.append(f"{issue.description}: {e}")
        # 保存文件
        save_path = filedialog.asksaveasfilename(
            title="保存修改后的文件",
            defaultextension=".docx",
            initialfile=Path(self.doc_path).stem + "_已修正.docx",
            filetypes=[("Word文档", "*.docx")])
        if save_path:
            try:
                self.doc.save(save_path)
                msg = f"已成功应用 {len(selected) - len(errors)} 项修改，保存至:\n{save_path}"
                if errors:
                    msg += f"\n\n{len(errors)} 项修改失败:\n" + "\n".join(errors[:5])
                messagebox.showinfo("完成", msg)
                # 重新检查
                self.doc = Document(save_path)
                self.doc_path = save_path
                self.file_label.config(text=Path(save_path).name)
                self._run_check()
            except Exception as e:
                messagebox.showerror("错误", f"保存失败: {e}")


def main():
    if len(sys.argv) > 1:
        # 命令行模式：直接检查并输出报告
        doc_path = sys.argv[1]
        # 解析可选的 --keywords 参数
        keywords = None
        if '--keywords' in sys.argv:
            idx = sys.argv.index('--keywords')
            if idx + 1 < len(sys.argv):
                keywords = [k.strip() for k in sys.argv[idx+1].split(',') if k.strip()]
        print(f"正在检查: {doc_path}")
        if keywords:
            print(f"身份关键词: {', '.join(keywords)}")
        try:
            doc = Document(doc_path)
        except Exception as e:
            print(f"错误: 无法打开文件 - {e}")
            sys.exit(1)
        issues = run_all_checks(doc, keywords)
        if not issues:
            print("✓ 所有检查通过，文档符合暗标编制要求！")
        else:
            print(f"\n共发现 {len(issues)} 个不符合项:\n")
            for i, issue in enumerate(issues, 1):
                fixable = " [可修复]" if issue.fix_func else ""
                print(f"  {i}. [{issue.category}] {issue.description}")
                print(f"     位置: {issue.location}{fixable}")
            print(f"\n提示: 运行不带参数启动GUI模式可交互式修复。")
            print(f"      使用 --keywords \"公司A,张三\" 检测身份信息。")
    else:
        # GUI模式
        root = tk.Tk()
        app = CheckerApp(root)
        root.mainloop()


if __name__ == "__main__":
    main()
