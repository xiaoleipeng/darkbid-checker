#!/usr/bin/env python3
"""
Word文档暗标编制要求检查工具 - GUI版
支持选择文件、勾选修复类别、自动修复+批注、展示修复结果。
需要 python-docx >= 1.2.0
"""

import sys
import re
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# 常量
FONT_NAME = '宋体'
FONT_SIZE = Pt(14)
PAGE_WIDTH = Cm(21.0)
PAGE_HEIGHT = Cm(29.7)
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.0)
MARGIN_LEFT = Cm(2.0)
MARGIN_RIGHT = Cm(2.0)
LINE_SPACING_VAL = Pt(30)
BLACK_COLOR = RGBColor(0, 0, 0)
AUTHOR = "暗标检查工具"
INITIALS = "AH"

CATEGORY_LABELS = {
    'page': '页面设置（A4纸张、上边距2.5cm其余2cm、清除页眉页脚页码、检测目录）',
    'spacing': '行间距固定值30磅、段前段后间距为0',
    'align': '左对齐、首行缩进2字符、删除开头空格',
    'font': '宋体四号（表格内不改字体字号）、去除加粗/倾斜/下划线（含表格内）',
    'color': '所有文字颜色→黑色（含表格内）',
    'punct': '全部使用中文标点（英文标点→中文标点）',
    'identity': '屏蔽投标人名称、人员姓名等身份信息',
}

PUNCTUATION_MAP = {
    ',': '，', '.': '。', ':': '：', ';': '；',
    '!': '！', '?': '？', '(': '（', ')': '）',
    '[': '【', ']': '】',
}
DIGIT_PUNCT_PATTERN = re.compile(r'\d[.,]\d')


def _is_in_table(paragraph):
    parent = paragraph._element.getparent()
    while parent is not None:
        if parent.tag.endswith('}tc') or parent.tag.endswith('}tbl'):
            return True
        parent = parent.getparent()
    return False


def _comment(doc, runs, text):
    doc.add_comment(runs=runs, text=text, author=AUTHOR, initials=INITIALS)


def _set_font_name(run):
    run.font.name = FONT_NAME
    rpr = run._element.find(qn('w:rPr'))
    if rpr is None:
        rpr = run._element.makeelement(qn('w:rPr'), {})
        run._element.insert(0, rpr)
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)


def _fix_punctuation_in_run(run):
    text = run.text
    result = []
    for i, ch in enumerate(text):
        if ch in PUNCTUATION_MAP:
            before_digit = i > 0 and text[i-1].isdigit()
            after_digit = i < len(text)-1 and text[i+1].isdigit()
            if before_digit and after_digit and ch in '.,':
                result.append(ch)
            else:
                result.append(PUNCTUATION_MAP[ch])
        else:
            result.append(ch)
    run.text = ''.join(result)


def check_fix_annotate(doc_path, output_path, keyword_map=None, categories=None, add_comments=True):
    """核心修复逻辑，返回 (修复详情列表, fix_count, warn_count)"""
    if categories is None:
        categories = set(CATEGORY_LABELS.keys())

    doc = Document(doc_path)
    fix_count = 0
    warn_count = 0
    details = []  # [(位置, 修复描述)]

    first_para = doc.paragraphs[0] if doc.paragraphs else None

    # === TOC目录检测（仅报告，无法自动修复） ===
    if 'page' in categories:
        for para in doc.paragraphs:
            for instr in para._element.findall('.//' + qn('w:instrText')):
                if instr.text and 'TOC' in instr.text.upper():
                    details.append(("目录", "⚠️ 存在自动目录，需手动删除（技术标暗标部分不得设置目录）"))
                    warn_count += 1
                    if add_comments and para.runs:
                        _comment(doc, para.runs, "【⚠️ 需手动处理】存在自动目录，技术标暗标部分不得设置目录，请手动删除。")
                    break

    # === 页数估算 ===
    total_paras = len(doc.paragraphs)
    estimated_pages = total_paras // 25
    if estimated_pages > 300:
        details.append(("全文", f"⚠️ 估算页数约{estimated_pages}页，可能超过300页限制"))
        warn_count += 1

    # === page ===
    if 'page' in categories:
        page_fixes = []
        tolerance = Cm(0.1)
        margin_tolerance = Cm(0.05)
        for section in doc.sections:
            if section.page_width and abs(section.page_width - PAGE_WIDTH) > tolerance:
                page_fixes.append(f"纸张宽度: {section.page_width.cm:.1f}cm → 21.0cm")
                section.page_width = PAGE_WIDTH
            if section.page_height and abs(section.page_height - PAGE_HEIGHT) > tolerance:
                page_fixes.append(f"纸张高度: {section.page_height.cm:.1f}cm → 29.7cm")
                section.page_height = PAGE_HEIGHT
            if section.top_margin is not None and abs(section.top_margin - MARGIN_TOP) > margin_tolerance:
                page_fixes.append(f"上边距: {section.top_margin.cm:.2f}cm → 2.5cm")
                section.top_margin = MARGIN_TOP
            if section.bottom_margin is not None and abs(section.bottom_margin - MARGIN_BOTTOM) > margin_tolerance:
                page_fixes.append(f"下边距: {section.bottom_margin.cm:.2f}cm → 2.0cm")
                section.bottom_margin = MARGIN_BOTTOM
            if section.left_margin is not None and abs(section.left_margin - MARGIN_LEFT) > margin_tolerance:
                page_fixes.append(f"左边距: {section.left_margin.cm:.2f}cm → 2.0cm")
                section.left_margin = MARGIN_LEFT
            if section.right_margin is not None and abs(section.right_margin - MARGIN_RIGHT) > margin_tolerance:
                page_fixes.append(f"右边距: {section.right_margin.cm:.2f}cm → 2.0cm")
                section.right_margin = MARGIN_RIGHT
            if section.header and section.header.is_linked_to_previous is False:
                header_text = ''.join(p.text for p in section.header.paragraphs).strip()
                if header_text:
                    page_fixes.append(f"已清除页眉")
                    for p in section.header.paragraphs:
                        p.clear()
            if section.footer and section.footer.is_linked_to_previous is False:
                footer_text = ''.join(p.text for p in section.footer.paragraphs).strip()
                if footer_text:
                    page_fixes.append("已清除页脚/页码")
                    for p in section.footer.paragraphs:
                        p.clear()
        if page_fixes and first_para and first_para.runs:
            if add_comments:
                _comment(doc, first_para.runs, "【已修复 - 页面设置】\n" + "\n".join(f"• {f}" for f in page_fixes))
            fix_count += len(page_fixes)
            for f in page_fixes:
                details.append(("页面设置", f))

    # === 逐段 ===
    for idx, para in enumerate(doc.paragraphs):
        if not para.text.strip() or not para.runs:
            continue
        in_table = _is_in_table(para)
        para_loc = para.text[:25] + "..." if len(para.text) > 25 else para.text
        para_fixes = []

        # 以下格式检查仅针对非表格段落
        if not in_table:
            if 'spacing' in categories:
                pf = para.paragraph_format
                if pf.line_spacing_rule != WD_LINE_SPACING.EXACTLY or pf.line_spacing != LINE_SPACING_VAL:
                    para_fixes.append("行间距 → 固定值30磅")
                    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    pf.line_spacing = LINE_SPACING_VAL
                if pf.space_before and pf.space_before > Pt(0):
                    para_fixes.append("段前间距 → 0")
                    pf.space_before = Pt(0)
                if pf.space_after and pf.space_after > Pt(0):
                    para_fixes.append("段后间距 → 0")
                    pf.space_after = Pt(0)

            if 'align' in categories:
                pf = para.paragraph_format
                if pf.alignment is not None and pf.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                    para_fixes.append("对齐 → 左对齐")
                    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if pf.first_line_indent is None or abs(pf.first_line_indent - Pt(28)) > Pt(1):
                    para_fixes.append("首行缩进 → 2字符")
                    pf.first_line_indent = Pt(28)
                if para.text.startswith(' ') or para.text.startswith('\u3000'):
                    para_fixes.append("删除开头空格")
                    para.runs[0].text = para.runs[0].text.lstrip(' \u3000')

            if 'punct' in categories:
                cleaned = DIGIT_PUNCT_PATTERN.sub('', para.text)
                en_puncts = set(ch for ch in cleaned if ch in PUNCTUATION_MAP)
                if en_puncts:
                    para_fixes.append(f"标点: {' '.join(repr(c) for c in sorted(en_puncts))} → 中文")
                    for run in para.runs:
                        _fix_punctuation_in_run(run)

        if 'identity' in categories and keyword_map:
            for kw, repl in keyword_map.items():
                if kw in para.text:
                    para_fixes.append(f"身份信息 '{kw}' → '{repl}'")
                    for run in para.runs:
                        if kw in run.text:
                            run.text = run.text.replace(kw, repl)
                    warn_count += 1

        if para_fixes:
            if add_comments:
                _comment(doc, para.runs, "【已修复】\n" + "\n".join(f"• {f}" for f in para_fixes))
            fix_count += len(para_fixes)
            for f in para_fixes:
                details.append((para_loc, f))

        # 字体检查：表格内只检查加粗/倾斜/下划线/颜色，不检查字体名和字号
        if 'font' in categories or 'color' in categories:
            for run in para.runs:
                if not run.text.strip():
                    continue
                font = run.font
                run_fixes = []
                run_loc = run.text[:15] + "..." if len(run.text) > 15 else run.text

                if 'font' in categories:
                    # 字体名和字号：仅非表格段落检查
                    if not in_table:
                        if font.name and font.name != FONT_NAME:
                            run_fixes.append(f"字体: {font.name} → 宋体")
                            _set_font_name(run)
                        else:
                            rpr = run._element.find(qn('w:rPr'))
                            if rpr is not None:
                                rFonts = rpr.find(qn('w:rFonts'))
                                if rFonts is not None:
                                    ea = rFonts.get(qn('w:eastAsia'))
                                    if ea and ea != FONT_NAME:
                                        run_fixes.append(f"中文字体: {ea} → 宋体")
                                        _set_font_name(run)
                        if font.size and font.size != FONT_SIZE:
                            run_fixes.append(f"字号: {font.size.pt:.1f}pt → 14pt")
                            font.size = FONT_SIZE
                    # 加粗/倾斜/下划线：所有文字都检查（含表格内）
                    if font.bold:
                        run_fixes.append("去除加粗")
                        font.bold = False
                    if font.italic:
                        run_fixes.append("去除倾斜")
                        font.italic = False
                    if font.underline:
                        run_fixes.append("去除下划线")
                        font.underline = False

                if 'color' in categories:
                    if font.color and font.color.rgb and font.color.rgb != BLACK_COLOR:
                        run_fixes.append(f"颜色: #{font.color.rgb} → 黑色")
                        font.color.rgb = BLACK_COLOR

                if run_fixes:
                    if add_comments:
                        prefix = "【已修复 - 表格内字体】" if in_table else "【已修复 - 字体】"
                        _comment(doc, [run], prefix + "\n" + "\n".join(f"• {f}" for f in run_fixes))
                    fix_count += len(run_fixes)
                    loc_prefix = "[表格内] " if in_table else ""
                    for f in run_fixes:
                        details.append((loc_prefix + run_loc, f))

    # === 表格内段落：检查对齐、加粗/倾斜/下划线/颜色（字体名和字号不检查） ===
    if 'font' in categories or 'color' in categories or 'align' in categories:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text.strip() or not para.runs:
                            continue
                        para_loc = "[表格内] " + (para.text[:15] + "..." if len(para.text) > 15 else para.text)

                        # 表格内对齐检查
                        if 'align' in categories:
                            pf = para.paragraph_format
                            if pf.alignment is not None and pf.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                                align_map = {WD_ALIGN_PARAGRAPH.CENTER: "居中", WD_ALIGN_PARAGRAPH.RIGHT: "右对齐", WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐"}
                                fix_desc = f"对齐: {align_map.get(pf.alignment, '?')} → 左对齐"
                                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                if add_comments:
                                    _comment(doc, para.runs, "【已修复 - 表格内】\n• " + fix_desc)
                                fix_count += 1
                                details.append((para_loc, fix_desc))

                        # 表格内字体样式检查
                        if 'font' in categories or 'color' in categories:
                            for run in para.runs:
                                if not run.text.strip():
                                    continue
                                font = run.font
                                run_fixes = []
                                run_loc = "[表格内] " + (run.text[:15] + "..." if len(run.text) > 15 else run.text)

                                if 'font' in categories:
                                    if font.bold:
                                        run_fixes.append("去除加粗")
                                        font.bold = False
                                    if font.italic:
                                        run_fixes.append("去除倾斜")
                                        font.italic = False
                                    if font.underline:
                                        run_fixes.append("去除下划线")
                                        font.underline = False

                                if 'color' in categories:
                                    if font.color and font.color.rgb and font.color.rgb != BLACK_COLOR:
                                        run_fixes.append(f"颜色: #{font.color.rgb} → 黑色")
                                        font.color.rgb = BLACK_COLOR

                                if run_fixes:
                                    if add_comments:
                                        _comment(doc, [run], "【已修复 - 表格内字体】\n" + "\n".join(f"• {f}" for f in run_fixes))
                                    fix_count += len(run_fixes)
                                    for f in run_fixes:
                                        details.append((run_loc, f))

    if fix_count == 0 and warn_count == 0:
        return details, 0, 0

    doc.save(output_path)
    return details, fix_count, warn_count


class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Word暗标编制要求检查修复工具")
        self.root.geometry("900x650")
        self.doc_path = None
        self.keyword_map = {}  # {关键词: 替换文本}
        self._build_ui()

    def _build_ui(self):
        # 文件选择
        f1 = tk.Frame(self.root)
        f1.pack(fill=tk.X, padx=10, pady=5)
        tk.Button(f1, text="选择Word文件", command=self._open_file).pack(side=tk.LEFT)
        self.file_label = tk.Label(f1, text="未选择文件", fg="gray")
        self.file_label.pack(side=tk.LEFT, padx=10)

        # 类别勾选
        f2 = tk.LabelFrame(self.root, text="选择修复类别")
        f2.pack(fill=tk.X, padx=10, pady=5)
        self.cat_vars = {}
        for i, (key, label) in enumerate(CATEGORY_LABELS.items()):
            var = tk.BooleanVar(value=True)
            self.cat_vars[key] = var
            tk.Checkbutton(f2, text=f"{key}: {label}", variable=var, anchor=tk.W).grid(row=i, column=0, sticky=tk.W, padx=5)

        # 关键词
        f3 = tk.Frame(self.root)
        f3.pack(fill=tk.X, padx=10, pady=3)
        tk.Label(f3, text="身份关键词:").pack(side=tk.LEFT)
        self.kw_entry = tk.Entry(f3, width=35)
        self.kw_entry.pack(side=tk.LEFT, padx=5)
        tk.Label(f3, text="替换为:").pack(side=tk.LEFT)
        self.replace_entry = tk.Entry(f3, width=15)
        self.replace_entry.insert(0, "***")
        self.replace_entry.pack(side=tk.LEFT, padx=5)
        tk.Button(f3, text="导入关键词文件", command=self._load_keywords).pack(side=tk.LEFT, padx=5)
        
        # 关键词提示
        f3b = tk.Frame(self.root)
        f3b.pack(fill=tk.X, padx=10)
        self.kw_info_label = tk.Label(f3b, text="（支持导入txt文件，格式每行: 关键词,替换文本）", fg="gray", font=("", 9))
        self.kw_info_label.pack(side=tk.LEFT)

        # 批注选项
        self.comment_var = tk.BooleanVar(value=True)
        tk.Checkbutton(f3, text="添加批注说明", variable=self.comment_var).pack(side=tk.LEFT, padx=10)

        # 执行按钮
        f4 = tk.Frame(self.root)
        f4.pack(fill=tk.X, padx=10, pady=5)
        tk.Button(f4, text="执行检查并修复", command=self._run, bg="#4CAF50", fg="white",
                  font=("", 11, "bold")).pack(side=tk.LEFT)
        self.open_dir_btn = tk.Button(f4, text="打开输出目录", command=self._open_output_dir, state=tk.DISABLED)
        self.open_dir_btn.pack(side=tk.LEFT, padx=10)
        self.status_label = tk.Label(f4, text="", fg="blue")
        self.status_label.pack(side=tk.LEFT, padx=10)
        self.output_path = None

        # 结果列表
        f5 = tk.LabelFrame(self.root, text="修复结果")
        f5.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        columns = ("location", "fix")
        self.tree = ttk.Treeview(f5, columns=columns, show="headings")
        self.tree.heading("location", text="位置")
        self.tree.heading("fix", text="修复内容")
        self.tree.column("location", width=250)
        self.tree.column("fix", width=550)

        sb = ttk.Scrollbar(f5, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=sb.set)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb.pack(side=tk.RIGHT, fill=tk.Y)

    def _open_file(self):
        path = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")])
        if path:
            self.doc_path = path
            self.file_label.config(text=Path(path).name, fg="black")

    def _open_output_dir(self):
        if self.output_path:
            import subprocess, platform
            folder = str(Path(self.output_path).parent)
            if platform.system() == 'Windows':
                subprocess.Popen(['explorer', '/select,', self.output_path])
            elif platform.system() == 'Darwin':
                subprocess.Popen(['open', folder])
            else:
                subprocess.Popen(['xdg-open', folder])

    def _load_keywords(self):
        """从文件加载关键词映射，格式每行: 关键词,替换文本"""
        path = filedialog.askopenfilename(
            title="选择关键词文件",
            filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")])
        if not path:
            return
        try:
            self.keyword_map = {}
            with open(path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith('#'):
                        continue
                    if ',' in line:
                        parts = line.split(',', 1)
                        self.keyword_map[parts[0].strip()] = parts[1].strip()
                    else:
                        self.keyword_map[line] = "***"
            self.kw_info_label.config(
                text=f"已导入 {len(self.keyword_map)} 条关键词映射（来自 {Path(path).name}）", fg="blue")
        except Exception as e:
            messagebox.showerror("错误", f"读取关键词文件失败: {e}")

    def _run(self):
        if not self.doc_path:
            messagebox.showwarning("提示", "请先选择Word文件")
            return

        categories = set(k for k, v in self.cat_vars.items() if v.get())
        if not categories:
            messagebox.showwarning("提示", "请至少选择一个修复类别")
            return

        kw_text = self.kw_entry.get().strip()
        replace_text = self.replace_entry.get() or "***"
        # 合并：文件导入的映射 + 手动输入的关键词
        keyword_map = dict(self.keyword_map)  # 复制文件导入的
        if kw_text:
            for k in kw_text.split(','):
                k = k.strip()
                if k and k not in keyword_map:
                    keyword_map[k] = replace_text

        p = Path(self.doc_path)
        if categories == set(CATEGORY_LABELS.keys()):
            suffix = '_已修正'
        else:
            suffix = '_' + '+'.join(sorted(categories))
        output_path = p.with_stem(p.stem + suffix).as_posix()

        self.status_label.config(text="正在处理...")
        self.root.update()

        try:
            details, fix_count, warn_count = check_fix_annotate(
                self.doc_path, output_path, keyword_map=keyword_map or None,
                categories=categories, add_comments=self.comment_var.get())
        except Exception as e:
            messagebox.showerror("错误", str(e))
            self.status_label.config(text="")
            return

        # 刷新结果列表
        self.tree.delete(*self.tree.get_children())

        if fix_count == 0 and warn_count == 0:
            self.status_label.config(text="✓ 所有检查通过，文档符合暗标编制要求！", fg="green")
            self.open_dir_btn.config(state=tk.DISABLED)
        else:
            self.status_label.config(
                text=f"已修复 {fix_count} 项，输出: {Path(output_path).name}", fg="blue")
            self.output_path = output_path
            self.open_dir_btn.config(state=tk.NORMAL)
            for loc, fix in details:
                self.tree.insert("", tk.END, values=(loc, fix))


def main():
    if len(sys.argv) > 1 and sys.argv[1] != '--gui':
        # CLI模式保留
        from check_word_comment import main as cli_main
        cli_main()
    else:
        root = tk.Tk()
        App(root)
        root.mainloop()


if __name__ == "__main__":
    main()
