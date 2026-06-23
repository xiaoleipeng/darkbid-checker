#!/usr/bin/env python3
"""
Word文档暗标编制要求检查工具 - 自动修复+批注模式
自动修复不符合项，同时在修改位置添加批注说明修改内容。
支持 --only 参数分类修复：page,spacing,align,font,color,punct,identity
需要 python-docx >= 1.2.0
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# 常量
FONT_NAME = '宋体'
FONT_SIZE = Pt(14)
PAGE_WIDTH = Cm(21.0)
PAGE_HEIGHT = Cm(29.7)
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.0)
MARGIN_LEFT = Cm(2.0)
MARGIN_RIGHT = Cm(2.0)
LINE_SPACING_VAL = Pt(30)
BLACK_COLOR = RGBColor(0, 0, 0)
AUTHOR = "暗标检查工具"
INITIALS = "AH"

ALL_CATEGORIES = {'page', 'spacing', 'align', 'font', 'color', 'punct', 'identity'}

PUNCTUATION_MAP = {
    ',': '，', '.': '。', ':': '：', ';': '；',
    '!': '！', '?': '？', '(': '（', ')': '）',
    '[': '【', ']': '】',
}
DIGIT_PUNCT_PATTERN = re.compile(r'\d[.,]\d')


def _is_in_table(paragraph):
    parent = paragraph._element.getparent()
    while parent is not None:
        if parent.tag.endswith('}tc') or parent.tag.endswith('}tbl'):
            return True
        parent = parent.getparent()
    return False


def _comment(doc, runs, text):
    doc.add_comment(runs=runs, text=text, author=AUTHOR, initials=INITIALS)


def _set_font_name(run):
    run.font.name = FONT_NAME
    rpr = run._element.find(qn('w:rPr'))
    if rpr is None:
        rpr = run._element.makeelement(qn('w:rPr'), {})
        run._element.insert(0, rpr)
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)


def _fix_punctuation_in_run(run):
    text = run.text
    result = []
    for i, ch in enumerate(text):
        if ch in PUNCTUATION_MAP:
            before_digit = i > 0 and text[i-1].isdigit()
            after_digit = i < len(text)-1 and text[i+1].isdigit()
            if before_digit and after_digit and ch in '.,':
                result.append(ch)
            else:
                result.append(PUNCTUATION_MAP[ch])
        else:
            result.append(ch)
    run.text = ''.join(result)


def check_fix_annotate(doc_path, output_path, keywords=None, categories=None):
    if categories is None:
        categories = ALL_CATEGORIES

    doc = Document(doc_path)
    fix_count = 0
    warn_count = 0

    first_para = doc.paragraphs[0] if doc.paragraphs else None

    # === 1. page: 页面设置 ===
    if 'page' in categories:
        page_fixes = []
        tolerance = Cm(0.1)
        margin_tolerance = Cm(0.05)
        for section in doc.sections:
            if section.page_width and abs(section.page_width - PAGE_WIDTH) > tolerance:
                page_fixes.append(f"纸张宽度: {section.page_width.cm:.1f}cm → 21.0cm")
                section.page_width = PAGE_WIDTH
            if section.page_height and abs(section.page_height - PAGE_HEIGHT) > tolerance:
                page_fixes.append(f"纸张高度: {section.page_height.cm:.1f}cm → 29.7cm")
                section.page_height = PAGE_HEIGHT
            if section.top_margin is not None and abs(section.top_margin - MARGIN_TOP) > margin_tolerance:
                page_fixes.append(f"上边距: {section.top_margin.cm:.2f}cm → 2.5cm")
                section.top_margin = MARGIN_TOP
            if section.bottom_margin is not None and abs(section.bottom_margin - MARGIN_BOTTOM) > margin_tolerance:
                page_fixes.append(f"下边距: {section.bottom_margin.cm:.2f}cm → 2.0cm")
                section.bottom_margin = MARGIN_BOTTOM
            if section.left_margin is not None and abs(section.left_margin - MARGIN_LEFT) > margin_tolerance:
                page_fixes.append(f"左边距: {section.left_margin.cm:.2f}cm → 2.0cm")
                section.left_margin = MARGIN_LEFT
            if section.right_margin is not None and abs(section.right_margin - MARGIN_RIGHT) > margin_tolerance:
                page_fixes.append(f"右边距: {section.right_margin.cm:.2f}cm → 2.0cm")
                section.right_margin = MARGIN_RIGHT
            if section.header and section.header.is_linked_to_previous is False:
                header_text = ''.join(p.text for p in section.header.paragraphs).strip()
                if header_text:
                    page_fixes.append(f"已清除页眉: '{header_text[:20]}'")
                    for p in section.header.paragraphs:
                        p.clear()
            if section.footer and section.footer.is_linked_to_previous is False:
                footer_text = ''.join(p.text for p in section.footer.paragraphs).strip()
                if footer_text:
                    page_fixes.append("已清除页脚/页码")
                    for p in section.footer.paragraphs:
                        p.clear()
        if page_fixes and first_para and first_para.runs:
            _comment(doc, first_para.runs, "【已修复 - 页面设置】\n" + "\n".join(f"• {f}" for f in page_fixes))
            fix_count += len(page_fixes)

    # === 2. 逐段检查 ===
    for para in doc.paragraphs:
        if not para.text.strip() or _is_in_table(para) or not para.runs:
            continue

        para_fixes = []

        # spacing: 行间距、段前段后
        if 'spacing' in categories:
            pf = para.paragraph_format
            if pf.line_spacing_rule != WD_LINE_SPACING.EXACTLY or pf.line_spacing != LINE_SPACING_VAL:
                para_fixes.append("行间距 → 固定值30磅")
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = LINE_SPACING_VAL
            if pf.space_before and pf.space_before > Pt(0):
                para_fixes.append("段前间距 → 0")
                pf.space_before = Pt(0)
            if pf.space_after and pf.space_after > Pt(0):
                para_fixes.append("段后间距 → 0")
                pf.space_after = Pt(0)

        # align: 对齐、首行缩进、开头空格
        if 'align' in categories:
            pf = para.paragraph_format
            if pf.alignment is not None and pf.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                align_map = {WD_ALIGN_PARAGRAPH.CENTER: "居中", WD_ALIGN_PARAGRAPH.RIGHT: "右对齐", WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐"}
                para_fixes.append(f"对齐: {align_map.get(pf.alignment, '?')} → 左对齐")
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if pf.first_line_indent is None or abs(pf.first_line_indent - Pt(28)) > Pt(1):
                para_fixes.append("首行缩进 → 2字符")
                pf.first_line_indent = Pt(28)
            if para.text.startswith(' ') or para.text.startswith('\u3000'):
                para_fixes.append("已删除段落开头空格")
                para.runs[0].text = para.runs[0].text.lstrip(' \u3000')

        # punct: 英文标点
        if 'punct' in categories:
            cleaned = DIGIT_PUNCT_PATTERN.sub('', para.text)
            en_puncts = set(ch for ch in cleaned if ch in PUNCTUATION_MAP)
            if en_puncts:
                para_fixes.append(f"英文标点 → 中文标点: {' '.join(repr(c) for c in sorted(en_puncts))}")
                for run in para.runs:
                    _fix_punctuation_in_run(run)

        # identity: 身份信息
        if 'identity' in categories and keywords:
            for kw in keywords:
                if kw in para.text:
                    para_fixes.append(f"身份信息 '{kw}' → ***")
                    for run in para.runs:
                        if kw in run.text:
                            run.text = run.text.replace(kw, '***')
                    warn_count += 1

        if para_fixes:
            _comment(doc, para.runs, "【已修复】\n" + "\n".join(f"• {f}" for f in para_fixes))
            fix_count += len(para_fixes)

        # === 3. font/color: 字体修复（精确到run）===
        if 'font' in categories or 'color' in categories:
            for run in para.runs:
                if not run.text.strip():
                    continue
                font = run.font
                run_fixes = []

                if 'font' in categories:
                    if font.name and font.name != FONT_NAME:
                        run_fixes.append(f"字体: {font.name} → 宋体")
                        _set_font_name(run)
                    else:
                        rpr = run._element.find(qn('w:rPr'))
                        if rpr is not None:
                            rFonts = rpr.find(qn('w:rFonts'))
                            if rFonts is not None:
                                ea = rFonts.get(qn('w:eastAsia'))
                                if ea and ea != FONT_NAME:
                                    run_fixes.append(f"中文字体: {ea} → 宋体")
                                    _set_font_name(run)
                    if font.size and font.size != FONT_SIZE:
                        run_fixes.append(f"字号: {font.size.pt:.1f}pt → 14pt(四号)")
                        font.size = FONT_SIZE
                    if font.bold:
                        run_fixes.append("已去除加粗")
                        font.bold = False
                    if font.italic:
                        run_fixes.append("已去除倾斜")
                        font.italic = False
                    if font.underline:
                        run_fixes.append("已去除下划线")
                        font.underline = False

                if 'color' in categories:
                    if font.color and font.color.rgb and font.color.rgb != BLACK_COLOR:
                        run_fixes.append(f"颜色: #{font.color.rgb} → 黑色")
                        font.color.rgb = BLACK_COLOR

                if run_fixes:
                    _comment(doc, [run], "【已修复 - 字体】\n" + "\n".join(f"• {f}" for f in run_fixes))
                    fix_count += len(run_fixes)

    if fix_count == 0 and warn_count == 0:
        return doc, 0, 0
    doc.save(output_path)
    return doc, fix_count, warn_count


def main():
    if len(sys.argv) < 2:
        print("用法: python3 check_word_comment.py <文件路径> [选项]")
        print()
        print("选项:")
        print("  --only <类别>    仅修复指定类别（逗号分隔）")
        print("  --keywords <词>  身份信息关键词（逗号分隔）")
        print()
        print("支持的类别:")
        print("  page      页面设置（纸张A4、页边距、页眉页脚）")
        print("  spacing   行间距（固定值30磅）、段前段后间距")
        print("  align     对齐方式（左对齐）、首行缩进、删除开头空格")
        print("  font      字体（宋体四号）、去除加粗/倾斜/下划线")
        print("  color     字体颜色（→黑色）")
        print("  punct     英文标点→中文标点")
        print("  identity  身份信息关键词替换（需配合--keywords）")
        print()
        print("示例:")
        print("  python3 check_word_comment.py 文件.docx")
        print("  python3 check_word_comment.py 文件.docx --only page,spacing")
        print("  python3 check_word_comment.py 文件.docx --only font,color")
        print("  python3 check_word_comment.py 文件.docx --keywords \"XX公司,张三\"")
        sys.exit(0)

    doc_path = sys.argv[1]
    keywords = None
    categories = None

    if '--keywords' in sys.argv:
        idx = sys.argv.index('--keywords')
        if idx + 1 < len(sys.argv):
            keywords = [k.strip() for k in sys.argv[idx+1].split(',') if k.strip()]

    if '--only' in sys.argv:
        idx = sys.argv.index('--only')
        if idx + 1 < len(sys.argv):
            categories = set(k.strip() for k in sys.argv[idx+1].split(',') if k.strip())
            invalid = categories - ALL_CATEGORIES
            if invalid:
                print(f"错误: 无效类别 {invalid}，支持: {', '.join(sorted(ALL_CATEGORIES))}")
                sys.exit(1)

    p = Path(doc_path)
    if categories:
        suffix = '_' + '+'.join(sorted(categories))
    else:
        suffix = '_已修正'
    output_path = p.with_stem(p.stem + suffix).as_posix()

    print(f"正在检查并修复: {doc_path}")
    if categories:
        print(f"修复类别: {', '.join(sorted(categories))}")
    if keywords:
        print(f"身份关键词: {', '.join(keywords)}")

    try:
        _, fix_count, warn_count = check_fix_annotate(doc_path, output_path, keywords, categories)
    except Exception as e:
        print(f"错误: {e}")
        import traceback; traceback.print_exc()
        sys.exit(1)

    if fix_count == 0 and warn_count == 0:
        print("✓ 所有检查通过，文档符合暗标编制要求！")
    else:
        print(f"\n已自动修复 {fix_count} 项问题" + (f"，{warn_count} 项警告" if warn_count else ""))
        print(f"输出文件: {output_path}")
        print("请在Word中打开，查看批注确认修改内容。")


if __name__ == "__main__":
    main()
